# 使用import导入openpyxl模块
import openpyxl
# 使用import导入docx
import docx

# 添加data_only = True，读取工作目录里名为"夜曲大学英语考试成绩.xlsx"的工作簿并赋值给变量wb
wb = openpyxl.load_workbook("夜曲大学英语考试成绩.xlsx",data_only = True)

# 通过工作簿对象wb获取名为“汇总”的工作表对象，并赋值给变量ws
ws = wb["汇总"]
# 设置空列表存储信息
nameList=[]
schoolList=[]
collegeList=[]
studentIDList=[]
certificateIDList=[]
listeningScoreList=[]
readingScoreList=[]
writingScoreList=[]
totalScoreList=[]

# 遍历工作表中每一行的数据.rows属性
for row in ws.rows:
    # 将这一行第1个单元格值，并赋值给变量name,再添加到列表中
    name = row[0].value
    nameList.append(name)
    # 将这一行第2个单元格值，并赋值给变量school，再添加到列表中
    school = row[1].value
    schoolList.append(school)
    # 将这一行第3个单元格值，并赋值给变量college，再添加到列表中
    college = row[2].value
    collegeList.append(college)
    # 将这一行第4个单元格值，并赋值给变量studentID，再添加到列表中
    studentID = row[3].value
    studentIDList.append(studentID)
    # 将这一行第5个单元格值，并赋值给变量certificateID，再添加到列表中
    certificateID = row[4].value
    certificateIDList.append(certificateID)
    # 将这一行第6个单元格值，用str()转换成字符串类型
    # 并赋值给变量listeningScore，再添加到列表中
    listeningScore = str(row[5].value)
    listeningScoreList.append(listeningScore)
    # 将这一行第7个单元格值，用str()转换成字符串类型
    # 并赋值给变量readingScore，再添加到列表中
    readingScore = str(row[6].value)
    readingScoreList.append(readingScore)   
    # 将这一行第8个单元格值，用str()转换成字符串类型
    # 并赋值给变量writingScore，再添加到列表中
    writingScore = str(row[7].value)
    writingScoreList.append(writingScore)
    # 将这一行第9个单元格值，用str()转换成字符串类型
    # 并赋值给变量totalScore，再添加到列表中
    totalScore = str(row[8].value)
    totalScoreList.append(totalScore)
 


for i in range(1,len(nameList)):
    # 读取工作目录里名为"成绩报告单模版.docx"的Word文件并赋值给变量doc
    doc = docx.Document("成绩报告单模版.docx")

    # 在Word文件的第4个段落的第3个样式块末尾，写入变量nameList[i]的文本内容
    doc.paragraphs[3].runs[2].add_text(nameList[i])

    # 在Word文件的第5个段落的第3个样式块末尾，写入变量schoolList[i]的文本内容
    doc.paragraphs[4].runs[2].add_text(schoolList[i])
    # 在Word文件的第6个段落的第2个样式块末尾，写入变量collegeList[i]的文本内容
    doc.paragraphs[5].runs[1].add_text(collegeList[i])
    # 在Word文件的第7个段落的第2个样式块末尾，写入变studentIDList[i]的文本内容
    doc.paragraphs[6].runs[1].add_text(studentIDList[i])
    # TODO 在Word文件的第11个段落的第1个样式块末尾，写入变量certificateIDList[i]的文本内容
    doc.paragraphs[10].runs[0].add_text(certificateIDList[i])
    # 读取Word文件的第1个表格，并赋值给变量table
    table = doc.tables[0]
    # 在表格的第2行第1列的空白单元格的段落内，写入变量totalScoreList[i]
    table.cell(1,0).paragraphs[0].add_run(totalScoreList[i])  
    # 在表格的第2行第2列的空白单元格的段落内，写入变量listeningScoreList[i]
    table.cell(1,1).paragraphs[0].add_run(listeningScoreList[i])
    # 在表格的第2行第3列的空白单元格的段落内，写入变量readingScoreList[i]
    table.cell(1,2).paragraphs[0].add_run(readingScoreList[i])
    # 在表格的第2行第4列的空白单元格的段落内，写入变量writingScoreList[i]
    table.cell(1,3).paragraphs[0].add_run(writingScoreList[i])

    # TODO 将Word文档以"学生成绩单/成绩报告单_{nameList[i]}.docx"的命名方式储存
    doc.save(f"学生成绩单/成绩报告单_{nameList[i]}.docx")
