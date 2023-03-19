# 使用import导入openpyxl模块
import openpyxl

# 使用import导入docx
import docx

# TODO 定义一个新函数replaceInfo，包含参数doc，oldInfo，newInfo
def replaceInfo(doc,oldInfo,newInfo):

    # TODO 遍历Word文件中的所有段落
    for paragraph in doc.paragraphs:

        # TODO 遍历所有段落中的所有样式块
        for run in paragraph.runs:

            # TODO 用replace()函数
            # 将样式块文本需要被替换的旧字符串oldInfo，替换成新字符串newInfo
            run.text = run.text.replace(oldInfo,newInfo)

    # TODO 遍历Word文件中的所有表格
    for table in doc.tables:
    
        # TODO 遍历所有表格中的所有行
        for row in table.rows:
        
            # TODO 遍历所有行中的所有单元格
            for cell in row.cells:
            
                # TODO 用replace()函数
                # 将单元格文本需要被替换的旧字符串oldInfo，替换成新字符串newInfo
                cell.text.replace(oldInfo,newInfo)

# 添加data_only = True，读取工作目录里名为"采购信息列表.xlsx"的工作簿并赋值给变量wb
wb = openpyxl.load_workbook("采购信息列表.xlsx",data_only = True)

# 通过工作簿对象wb获取名为“6月采购”的工作表对象，并赋值给变量ws
ws = wb["6月采购"]

# 建立一个空列表变量titleRow
titleRow = []

# TODO 使用enumerate()函数，遍历ws.rows的所有变量row，和对应行数索引rowIndex
for rowIndex,row in enumerate(ws.rows):

    # 读取工作目录里名为"合同模版.docx"的Word文件并赋值给doc
    doc = docx.Document("合同模版.docx")    
    
    # 判断如果行数rowIndex是0的话（也就是标题行）
    if rowIndex == 0:
        
        # TODO 将这一行赋值给变量titleRow
        titleRow = row
        
    # 否则
    else:
        
        # TODO 使用enumerate()函数，遍历这一行row的所有单元格cell，和对应列数索引columnIndex
        for columnIndex,cell in enumerate(row):

            # TODO 取标题行titleRow第columnIndex个元素的单元格值，赋值给oldInfo
            oldInfo = titleRow[columnIndex].value

            # TODO 取此时变量cell的单元格值
            # 使用str()函数转变成字符串格式，赋值给newInfo
            newInfo = str(cell.value)

            # TODO 使用replaceInfo()，对doc对象，将oldInfo的字符串内容，替换成newInfo的新字符串内容
            replaceInfo(doc,oldInfo,newInfo)

        # TODO 取这一行第4个元素的单元格值（也就是货物名称），赋值给变量productName
        productName = row[3].value

        # 写完一个合同，生成一个在“合同”文件夹下的Word文档
        doc.save(f"合同/采购合同_{productName}.docx")

# 如果运行成功，输出success
print("success")