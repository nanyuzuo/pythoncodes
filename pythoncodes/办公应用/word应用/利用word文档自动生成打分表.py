# 使用import导入os模块
import os
# 使用import导入docx
import docx
# 使用import导入openpyxl模块
import openpyxl

# 将统计表的路径 /Users/yequ/统计表 赋值给变量allFormPath
allFormPath = "/Users/yequ/统计表"
# 用os.listdir()函数获取该路径下所有的统计表，并赋值给变量allItems
allItems = os.listdir(allFormPath)

# TODO 定义一个空列表allStudentsData存储所有学生数据
allStudentsData = []

# 使用for循环遍历所有文件
for item in allItems:
    # TODO 定义一个空字典studentData存储单个学生数据
    studentData = {}

    # 使用os.path.join()函数拼接出统计表的路径，并赋值给变量formPath
    formPath = os.path.join(allFormPath, item)

    # TODO 使用docx.Document()函数读取统计表，并赋值给变量doc 
    doc = docx.Document(formPath)

    # TODO 读取统计表中的第一个表格，并赋值给变量table
    table = doc.tables[0]

    # TODO 将第三行“我是健康小卫士”这项中的第3个单元格内的学生姓名赋值给变量stuName
    stuName = table.cell(2,2).text

    # TODO 使用if语句判断当学生姓名为空时，即表明没有学生参加
    if stuName == "":

        # TODO 如果没有学生参加，就跳过当前统计表，遍历下一个
        continue

    # TODO 将学生姓名赋值到学生数据字典的studentName键里
    studentData["studentName"] = stuName

    # TODO 读取统计表中第二段的第二个样式块的文本内容，也就是学校名称，并赋值给变量schoolName
    schoolName = doc.paragraphs[1].runs[1].text

    # TODO 将学校名称schoolName赋值到学生数据字典的studentSchool键里
    studentData["studentSchool"] = schoolName
    
    # 将第三行“我是健康小卫士”这项中的第4个单元格内的学生班级
    # 赋值到学生数据字典的studentGrade键里
    studentData["studentGrade"] = table.cell(2,3).text
    
    # TODO 将第三行“我是健康小卫士”这项中的第5个单元格内的学生教师姓名
    # 赋值到学生数据字典的studentTeacher键里
    studentData["studentTeacher"] = table.cell(2,4).text

    # TODO 使用append()函数将studentData添加到总学生数据allStudentsData中
    allStudentsData.append(studentData)

# TODO 使用openpyxl.load_workbook()函数打开 /Users/yequ/我是健康小卫士打分表.xlsx，并赋值给变量wb
wb = openpyxl.load_workbook("/Users/yequ/我是健康小卫士打分表.xlsx")

# TODO 将名为 我是健康小卫士 的工作表赋值给变量sheet
sheet = wb["我是健康小卫士"]

# TODO 使用for循环和enumerate()函数
# 遍历所有学生数据allStudentsData的同时
# 生成一个从2开始的索引index
# 因为工作表中的第一行是表头并且在Excel中行数从1开始，所以我们的信息要从表格中的第二行开始填写
for index,studentData in enumerate(allStudentsData,2):

    # 给第A列数据设置学校名称
    sheet[f"A{index}"].value = studentData["studentSchool"]
    # TODO 给第B列数据设置学生姓名
    sheet[f"B{index}"].value = studentData["studentName"]

    # TODO 给第C列数据设置学生班级
    sheet[f"C{index}"].value = studentData["studentGrade"]

    # 给第D列数据设置学生的教师姓名
    sheet[f"D{index}"].value = studentData["studentTeacher"]

# TODO 使用save()函数将工作簿保存到指定路径 /Users/yequ/我是健康小卫士打分表.xlsx
wb.save("/Users/yequ/我是健康小卫士打分表.xlsx")