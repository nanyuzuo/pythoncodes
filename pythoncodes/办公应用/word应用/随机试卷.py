# 从"题库.docx"的5分题中随机抽取6道，从10分题中随机抽取4道，组成一张10道题的考卷，填入到"试卷模版.docx"中。然后生成72份这样随机的试卷

# 使用import导入docx
import docx

# 使用import导入random模块
import random

# 读取工作目录下"题库.docx"，赋值给变量questionBank
questionBank = docx.Document("题库.docx")

# 定义两个空列表questionList1、questionList2，用来储存读取出的题目
questionList1 = []
questionList2 = []


# 使用range()函数，用变量number，从0到11循环遍历
for number in range(0,12):
    
    # 读取题库文档questionBank的第number个段落的文本，赋值给变量title
    title = questionBank.paragraphs[number].text
    
    # 使用append()，将变量title的内容加入到列表questionList1中
    questionList1.append(title)
    
    
# 使用range()函数，用变量number，从14到21循环遍历
for number in range(14,22):
    
    # 读取题库文档questionBank的第number个段落的文本，赋值给变量title
    title = questionBank.paragraphs[number].text

    # 使用append()，将变量title的内容加入到列表中questionList2中
    questionList2.append(title)
    
    
# 使用range() 函数，从0到71循环遍历
for count in range(0,72):
    
    # 使用random.sample()函数，从列表questionList1的12个题目中
    # 随机抽取6个题目组成一个新的列表，赋值给formalQuestionList1
    formalQuestionList1 = random.sample(questionList1,6)

    # 使用random.sample()函数，从列表questionList2的8个题目中
    # 随机抽取4个题目组成一个新的列表，赋值给formalQuestionList2
    formalQuestionList2 = random.sample(questionList2,4)

    # 使用合并列表，将formalQuestionList1 和 formalQuestionList2 用"+"号连接，赋值给新的列表examQuestionList
    examQuestionList = formalQuestionList1 + formalQuestionList2

    # 读取工作目录下"试卷模版.docx"，赋值给变量examination
    examination = docx.Document("试卷模版.docx")

    # 使用enumerate()函数，遍历examQuestionList的所有变量examTitle，和对应索引index
    for index,examTitle in enumerate(examQuestionList):

        # 使用add_run()函数，将变量examTitle的内容，写入"试卷模版.docx"，第index段落末尾
        examination.paragraphs[index].add_run(examTitle)

    # 将生成的"试卷模版.docx"，命名为"试卷_{count+1}.docx"，保存在工作目录下"examination"文件夹中
    examination.save(f"examination/试卷_{count+1}.docx")

# 如果运行成功，输出"success"
print("success")