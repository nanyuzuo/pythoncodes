# 使用import导入openpyxl
import openpyxl

# 使用import导入docx
import docx

# TODO 定义一个新函数replaceInfo，包含参数doc，oldInfo，newInfo
def replaceInfo(doc,oldInfo,newInfo):
    # TODO 遍历Word文件中的所有段落
    for paragraph in doc.paragraphs:
        # TODO 遍历所有段落中的所有样式块
        for run in paragraph.runs:
            # TODO 用.replace()方法
            # 将样式块文本需要被替换的旧字符串oldInfo
            # 替换成新字符串newInfo
            run.text = run.text.replace(oldInfo,newInfo)
    # TODO 遍历Word文件中的所有表格
    for table in doc.tables:
        # TODO 遍历所有表格中的所有行
        for row in table.rows:
            # TODO 遍历所有行中的所有单元格
            for cell in row.cells:
                # TODO 用.replace()方法
                # 将单元格文本需要被替换的旧字符串oldInfo
                # 替换成新字符串newInfo
                cell.text = cell.text.replace(oldInfo,newInfo)

# 添加data_only = True，读取工作目录里名为"夜曲大学英语考试成绩.xlsx"的工作簿并赋值给变量wb
wb = openpyxl.load_workbook("夜曲大学英语考试成绩.xlsx",data_only = True)

# 通过工作簿对象wb获取名为“汇总”的工作表对象，并赋值给变量ws
ws = wb["汇总"]

# TODO 为表头定义一个空的元组并赋值给变量firstRow
firstRow = ()

# TODO 遍历工作表的所有行和其对应的索引
# 用变量rowIndex表示索引，变量row表示每一行
for rowIndex,row in enumerate(ws.rows):
    # 判断是第一行时
    if rowIndex == 0:
        # 将第一行数据赋值给变量firstRow
        firstRow = row
    # TODO 非第一行数据时
    else:
        # 读取模板Word文档：成绩报告单模版.docx并赋值给变量doc
        doc = docx.Document("成绩报告单模版.docx")
        # TODO 遍历每个单元格和对应的列索引
        # 用变量columnIndex表示索引，cell表示单元格
        for columnIndex,cell in enumerate(row):
            # TODO 将该单元格对应列的表头使用columnIndex提取出
            # 并赋值给变量oldInfo
            oldInfo = firstRow[columnIndex].value
            # TODO 将单元格的值使用str()函数转换成字符串
            # 并赋值给变量newInfo
            newInfo = str(cell.value)
            # TODO 使用replaceInfo()函数替换doc文档内容
            replaceInfo(doc,oldInfo,newInfo)
        # TODO 提取每行第一列的学生姓名，并赋值给变量name
        name = row[0].value
        # 将doc文档保存在 学生成绩单 文件夹下，文件名格式为"成绩报告单_{姓名}.docx"
        doc.save(f"学生成绩单/成绩报告单_{name}.docx")