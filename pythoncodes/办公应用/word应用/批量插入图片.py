# 导入docx
import docx
# 导入openpyxl模块
import openpyxl
# 导入os模块
import os

# 添加data_only = True，读取工作目录里名为"志愿者名单.xlsx"的工作簿并赋值给变量wb
wb = openpyxl.load_workbook("志愿者名单.xlsx",data_only = True)

# 读取名为"名单"的工作表并赋值给变量ws
ws = wb["名单"]

# 将照片的文件夹相对路径"photo"，赋值给变量photoPath
photoPath = "photo"

# TODO 使用for循环，遍历Excel文档的每一行
for row in ws.rows:

    # TODO 将这一行第1个元素的值，赋值给变量name
    name = row[0].value

    # TODO 将这一行第2个元素的值，赋值给变量group
    group = row[1].value

    # TODO 使用格式化输出，在name后加上.png，赋值给变量item
    item = f"{name}.png"
    
    # TODO 使用os.path.join()函数拼接文件路径photoPath和item，赋值给变量itemPath
    itemPath = os.path.join(photoPath,item)
    
    # 判断如果变量name是"姓名"的话
    if name == "姓名":
        
        # 就使用continue跳过，不执行写入
        continue
    
    # 读取工作目录下名为"志愿者工牌模版.docx"并赋值给变量doc
    doc = docx.Document("志愿者工牌模版.docx")

    # TODO 在Word文档的第8个段落第1个样式块末尾，添加name的文本
    doc.paragraphs[7].runs[0].add_text(name)

    # TODO 在Word文档的第9个段落第1个样式块末尾，添加group的文本 
    doc.paragraphs[8].runs[0].add_text(group)

    # TODO 在Word表格的空单元格中（即第1个表格第1行第1列的单元格第1个段落内）
    # 使用add_run()添加一个空的样式块，赋值给变量run
    run = doc.tables[0].cell(0,0).paragraphs[0].add_run()
    
    # TODO 在run对象上使用add_picture()函数，图片路径为itemPath，宽度使用docx.shared.Cm()函数限定为5
    run.add_picture(itemPath,width=docx.shared.Cm(5))

    # 完成一个工牌，生成一个在"志愿者"文件夹下的Word文档
    # 命名为"工牌_{name}.docx"
    doc.save(f"志愿者\\工牌_{name}.docx")

# 如果运行成功，输出success
print("success")