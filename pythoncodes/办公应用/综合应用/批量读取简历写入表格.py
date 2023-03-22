# 使用import导入os模块
import os
# 使用import导入docx
import docx
# 使用import导入openpyxl模块
import openpyxl

# 获取文件存储路径，将路径赋值给path
path = "/Users/ahua/teacher"
# 使用os.listdir()函数获取该路径下的所有文件并赋值给fileNames
fileNames = os.listdir(path)

# 获取汇总Excel表的存储路径,赋值给wb
wb = openpyxl.load_workbook("/Users/ahua/汇总.xlsx")

# 读取"教师信息登记表"，并赋值给aSheet
aSheet = wb["教师信息登记表"]

# 设置一个用于记录写入行数的count，初始值为2，从Excel的第二行写入
count = 2

# for循环遍历列表中的所有文件
for item in fileNames:

    # 使用os.path.splitext()函数获取文件后缀名,赋值给extension
    extension = os.path.splitext(item)[1]

    # TODO 判断后缀名不是 .docx时，就跳过
    if extension != ".docx":
        continue
        

    # 用os.path.join()函数拼接该文件的路径并赋值给filePath
    filePath = os.path.join(path, item)

    # 使用docx.Document()打开Word文档，赋值给teacherFile
    teacherFile = docx.Document(filePath)

    # TODO 使用.tables 使用.tables获取文档的第一个表格，并赋值给table
    table = teacherFile.tables[0]

    # 循环读取教师基本信息
    # TODO 读取第一行第二列的姓名，赋值给name
    name = table.cell(0,1).text
    gender = table.cell(0, 3).text
    birthDate = table.cell(0, 5).text
    highSchool = table.cell(3, 1).text
    masterSchool = table.cell(4, 1).text
    doctorSchool = table.cell(5, 1).text
    bestGraduate = table.cell(5, 5).text
    teachTime = table.cell(2, 1).text
    phoneNumber = table.cell(6, 1).text
    email = table.cell(6, 5).text
    job = table.cell(2, 5).text
    
    # 将读到的信息逐行写入单元格
    # TODO 将姓名赋值给Excel表中A列对应位置
    aSheet[f"A{count}"].value = name
    aSheet[f"B{count}"].value = gender
    aSheet[f"C{count}"].value = birthDate
    aSheet[f"D{count}"].value = highSchool
    aSheet[f"E{count}"].value = masterSchool
    aSheet[f"F{count}"].value = doctorSchool
    aSheet[f"G{count}"].value = bestGraduate
    aSheet[f"H{count}"].value = teachTime
    aSheet[f"I{count}"].value = job
    aSheet[f"J{count}"].value = phoneNumber
    aSheet[f"K{count}"].value = email

    # 格式化输出{item}写入完成
    print(f"{item}写入完成")

    # 写完一个count计数加1，写下一行
    count += 1

# 将Excel文档保存到路径"/Users/ahua/汇总.xlsx"
wb.save("/Users/ahua/汇总.xlsx")