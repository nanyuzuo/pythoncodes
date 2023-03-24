# 每个月的优秀稿件都会放进report文件夹中，但文档命名方式不统一。
# 者文章全部存储在：/Users/report/article路径下，每篇文章中第3段第1个样式块是姓名，第4个样式块是邮箱。
# 投稿Excel文档：第2列为姓名，第4列为邮箱
# 现在需要去“/Users/report/作者投稿-20年.xlsx”文件中找到文章收稿月份
# 然后按照月份重新命名文件“2020年1月小夜投稿.docx”
# 修改后存到"/Users/report/modify"路径下

# TODO 使用import导入os模块
import os

# TODO 使用import导入docx
import docx

# TODO 使用import导入openpyxl模块
import openpyxl

# TODO 使用openpyxl.load_workbook()打开"/Users/report/作者投稿-20年.xlsx"文档，赋值给wb
wb = openpyxl.load_workbook("/Users/report/作者投稿-20年.xlsx")

# TODO 将路径"/Users/report/article"，赋值给path
path = "/Users/report/article"

# TODO 使用os.listdir()获取路径下的所有文件名称，赋值给articleNames
articleNames = os.listdir(path)

# TODO for循环遍历所有的文件名称
for article in articleNames:

    # TODO 使用os.path.splitext()获取文件后缀名，赋值给extension
    extension = os.path.splitext(article)[1]

    # TODO 使用os.path.splitext()获取文件名称，赋值给name
    name = os.path.splitext(article)[0]

    # 使用if判断文件后缀名不是.docx就跳过
    if extension != '.docx':
        continue

    # TODO 使用os.path.join()函数拼接文件路径,赋值给artPath
    artPath = os.path.join(path,article)

    # TODO 使用Document获取docx文档,赋值给doc
    doc = docx.Document(artPath)

    # TODO 使用 .paragraphs 获取第3段，.runs 获取第1个样式块的内容，赋值给wName 
    wName = doc.paragraphs[2].runs[0].text

    # TODO 使用 .paragraphs 获取第3段，.runs 获取第4个样式块的内容，赋值给wEmail
    wEmail = doc.paragraphs[2].runs[3].text

    # TODO for和range(1, 7)逐个读取每个月份（共6个月）中的姓名和邮箱
    for i in range(1,7):

        # TODO 使用格式化生成每个月份的工作表名称，并获取工作表对象，赋值给monthSheet
        monthSheet = wb[f"{i}月"]

        # TODO for逐个读取每个月份工作表中的行
        for row in monthSheet.rows:

            # TODO 找到第2个索引为1的姓名的值，赋值给eName
            eName = row[1].value

            # TODO 找到第4个索引为3的邮箱的值，赋值给eEmail
            eEmail = row[3].value

            # TODO 使用if语句判断word文档中的姓名和邮箱和Excel文档中的姓名和邮箱均相等时
            if ( eName == wName ) and ( eEmail == wEmail ):

                # TODO 找到之后获取该工作表的名称，赋值给month
                month = monthSheet.title

                # 使用格式化输出 2020年x月份xx投稿.docx，并保存到 /Users/report/modify
                doc.save(f"/Users/report/modify/2020年{month}{name}.docx")

                # 使用print输出修改完成信息
                print(f"{name}修改完成")