import os
import docx
import openpyxl
import pdfplumber
import re

pattern = r"1\d{10}"

phoneList = []

def findPhonenumber(path):
    allItems = os.listdir(path)
    for item in allItems:
        filePath = os.path.join(path,item)
        if os.path.isdir(filePath):
            findPhonenumber(filePath)
        elif os.path.splitext(item)[1].lower() == ".docx":
            doc = docx.Document(filePath)
            for paragraph in doc.paragraphs:
                docnumber1 = re.findall(pattern,paragraph.text)
                phoneList.extend(docnumber1)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        docnumber2 = re.findall(pattern,cell.text)
                        phoneList.extend(docnumber2)
        elif os.path.splitext(item)[1].lower() == ".xlsx":
            wb = openpyxl.load_workbook(filePath)
            for sheet in wb.worksheets:
                for row in sheet.rows:
                    for cell in row:
                        exnumber = re.findall(pattern,str(cell.value))
                        phoneList.extend(exnumber)
        elif os.path.splitext(item)[1].lower() == ".pdf":
            pdf = pdfplumber.open(filePath)
            for page in pdf.pages:
                pdfnumber = re.findall(pattern,page.extract_text())
                phoneList.extend(pdfnumber)
        elif os.path.splitext(item)[1].lower() == ".txt":
            f = open(filePath,'r',encoding='UTF-8')
            txt = f.read()
            txtnumber = re.findall(pattern,txt)
            phoneList.extend(txtnumber)

findPhonenumber(r"d:\yequ\misc")
phoneList = list(set(phoneList))
print(phoneList)
newDoc = docx.Document()
numbers = " ".join(phoneList)
newDoc.add_paragraph(numbers)
newDoc.save(r"d:\yequ\misc\number.docx")
            