# TODO 使用import导入docx
import docx

# TODO 使用import导入os模块
import os

# TODO 获取需要修改的文件路径 /Users/records/stu ，赋值给path
path = "/Users/records/stu"

# TODO 获取该列表下的所有文件，赋值给fileNames
fileNames = os.listdir(path)

# TODO 循环遍历所有的文件
for file in fileNames:
    # TODO 使用os.path.splitext()获取文件后缀名，赋值给extension
    extension = os.path.splitext(file)[1]

    # 判断后缀名不是.docx文档就继续读下一个
    if extension != '.docx':
        continue

    # TODO 使用os.path.join()拼接待修改文档路径，赋值给filePath
    filePath = os.path.join(path,file)

    # TODO 使用docx.Document()获取文档，赋值给doc
    doc = docx.Document(filePath)

    # TODO 将“夜曲大学研究生院”赋值给第8段的文本
    doc.paragraphs[7].text = "夜曲大学研究生院"

    # TODO 使用原文档名称保存文档到 /Users/records/modify 路径下
    doc.save(f"/Users/records/modify/{file}")