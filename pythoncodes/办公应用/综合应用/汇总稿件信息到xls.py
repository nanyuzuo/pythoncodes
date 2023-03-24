# 作者投稿全部下载到了sub文件夹中，现在需要将文章的题目、作者、地址和邮箱复制粘贴到作者投稿.xlsx文档中的“6月”工作表中，文章中的姓名、地址和邮箱以空格分割。

# 文档路径：/Users/year/sub

# 整理文档路径：/Users/year/作者投稿.xlsx

# TODO 使用import导入os模块
import os

# TODO 使用import导入docx
import docx

# TODO 使用import导入openpyxl模块
import openpyxl

# TODO openpyxl.load_workbook()函数读取"/Users/year/作者投稿.xlsx"工作簿并赋值给wb 
wb = openpyxl.load_workbook("/Users/year/作者投稿.xlsx")

# TODO 通过工作簿对象wb获取名为“6月”的工作表对象，并赋值给变量june
june = wb["6月"]

# TODO 设置一个用于记录写入行数的counter，初始值为2，从Excel的第二行写入
count = 2

# TODO 将作者投稿文件夹路径 /Users/year/sub 赋值给变量path
path = "/Users/year/sub"

# TODO 使用os.listdir()函数获取该路径下所有的文件(夹)，并赋值给变量fileNames
fileNames = os.listdir(path)

# TODO for循环遍历列表中的所有文件
for file in fileNames:

    # TODO 使用os.path.join()函数拼接出投稿的路径，并赋值给变量filePath
    filePath = os.path.join(path,file)

    # TODO 使用docx.Document()读取投稿文档并赋值给变量doc 
    doc = docx.Document(filePath)

    # TODO 使用.paragraphs读取第一个段落的内容赋值给title
    title = doc.paragraphs[0].text

    # TODO 使用.paragraphs读取第3个段落的内容赋值给info
    info = doc.paragraphs[2].text

    # TODO 使用split()空格分割第3个段落中的名字、地址和邮箱，赋值给information
    information = info.split(" ")

    # TODO 将information的第一个元素给name
    name = information[0]
    
    # TODO 将information的第二个元素给address
    address = information[1]
    
    # TODO 将information的第三个元素给email
    email = information[2]
    
    # TODO 将读取的title赋值给6月工作表的A列单元格
    june[f"A{count}"].value = title
    
    # TODO 将读取name赋值给6月工作表的B列单元格
    june[f"B{count}"].value = name
    
    # TODO 将读取的address赋值给6月工作表的C列单元格
    june[f"C{count}"].value = address
    
    # TODO 将读取的email赋值给6月工作表的D列单元格
    june[f"D{count}"].value = email

    # 写完一行计数器加1写下一行
    count += 1

# TODO 将Excel文档保存至"/Users/year/作者投稿.xlsx"路径下
wb.save("/Users/year/作者投稿.xlsx")