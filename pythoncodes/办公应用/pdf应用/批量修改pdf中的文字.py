# 使用import导入os模块
import os

# TODO 使用import导入pdfplumber模块
import pdfplumber

# TODO 使用import导入docx
import docx

# TODO 新建一个空白的word文档赋值给变量file
file = docx.Document()

# 将PDF文件夹路径 /Users/xiaoqian/Desktop/文稿汇总，赋值给变量path
path = r"d:\yequ\文稿汇总"

# 使用os.listdir()函数获取该路径下所有的文件名称，并赋值给变量allItems
allItems = os.listdir(path)

# 定义一个nameList空列表
nameList = []

# 使用for循环遍历文件列表
for name in allItems:

    # 分割文件名，将文件名称添加到nameList中
    nameList.append(os.path.splitext(name)[0])

# 将nameList中的元素进行升序
nameList.sort()

# 定义一个newList空列表
newList = []

# 使用for循环遍历nameList
for i in nameList:

    # 所有元素后添加".pdf"，添加到newList列表中
    newList.append(i+".pdf")

# TODO 使用for循环遍历newList
for item in newList:

    # TODO 使用os.path.join()函数拼接pdf文件路径赋值给变量pdfPath
    pdfPath = os.path.join(path,item)

    # TODO 读取pdfPath中的文件，并赋值给pdf
    pdf = pdfplumber.open(pdfPath)

    # TODO 使用for循环遍历.pages属性，也就是所有页面对象
    for page in pdf.pages:

        # TODO 使用extract_text()函数提取页面中的文本,赋值给textData
        textData = page.extract_text()

        # TODO 替换python为Python
        textData = textData.replace("python","Python")

        # TODO 在Word文档file中使用add_paragraph()函数添加替换后的内容段落
        file.add_paragraph(textData)

    # TODO 格式化输出XX提取完成
    print(f"{item}提取完成")
        
    # TODO 添加分页符
    file.add_page_break()

# TODO 保存文档到指定路径/Users/xiaoqian/Desktop/Python合集.docx
file.save(r"d:\yequ\文稿汇总\Python合集.docx")